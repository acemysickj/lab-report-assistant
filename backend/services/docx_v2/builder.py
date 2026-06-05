"""python-docx builder helpers — extracted from verified reference scripts.

These functions are the stable API that AI-generated scripts call to build DOCX.
"""
from __future__ import annotations
import json
import re as _re_builtin
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from lxml import etree


# ── equation counter ────────────────────────────────────────────────────

class _EqCounter:
    def __init__(self):
        self.n = 0
    def next(self) -> int:
        self.n += 1
        return self.n

_eq = _EqCounter()


# ── low-level helpers ────────────────────────────────────────────────────

def clear_first_para(cell):
    """Remove default empty paragraph from a table cell."""
    if len(cell.paragraphs) == 1 and cell.paragraphs[0].text.strip() == '':
        p = cell.paragraphs[0]._element
        p.getparent().remove(p)


def set_cell_border(cell, **kw):
    """Set per-edge borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tb = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kw.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val","single")}" '
            f'w:sz="{attrs.get("sz",4)}" w:space="0" w:color="{attrs.get("color","000000")}"/>'
        )
        tb.append(el)
    tcPr.append(tb)


def set_tbl_borders(table, sz=4):
    """Add single-line borders (all edges, sz eighths-of-point) to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="000000"/>'
        f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    ))


def add_run(para, text, bold=False, size=Pt(12), ea='宋体'):
    """Add a text run with CJK font support."""
    r = para.add_run(text)
    r.bold = bold
    r.font.size = size
    rPr = r._element.get_or_add_rPr()
    rf = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{ea}"/>')
    ex = rPr.find(qn('w:rFonts'))
    if ex is not None:
        rPr.remove(ex)
    rPr.insert(0, rf)
    return r


# ── content helpers ──────────────────────────────────────────────────────

def heading(cell, text):
    """Section heading: 宋体四号 14pt bold.  Use for 一、二、三、四."""
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    add_run(p, text, bold=True, size=Pt(14))


def sub(cell, text):
    """Sub-heading: 宋体小四 12pt, NOT bold.  Use for 1. 2. 3."""
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    add_run(p, text, bold=False, size=Pt(12))


def body(cell, text, space_before=0):
    """Body text: 宋体小四 12pt, 1.5x line spacing."""
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(0)
    add_run(p, text, bold=False, size=Pt(12))


def body_lbl(cell, label, text):
    """Labeled body text: label + text on same line."""
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    add_run(p, label, bold=False, size=Pt(12))
    add_run(p, text, bold=False, size=Pt(12))


def body_sub(cell, text, space_before=0, formula_config_path=None):
    """Body text with auto-subscript rules and inline OMML ($...$).

    Reads subscript_rules from formula_config.json if available.
    """
    import re

    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(0)

    rules = []
    exceptions = set()
    if formula_config_path and Path(formula_config_path).exists():
        with open(formula_config_path, 'r', encoding='utf-8') as f:
            cfg = json.load(f)
        rules = cfg.get('subscript_rules', [])
        exceptions = set(cfg.get('subscript_exceptions', []))

    processed = text
    for rule in rules:
        pat = rule['pattern']
        def apply_sub(m, exc=exceptions):
            full = m.group(0)
            if full in exc:
                return full
            groups = m.groups()
            if len(groups) >= 2:
                return groups[0] + '\x00' + groups[1] + '\x01'
            return full
        processed = re.sub(pat, apply_sub, processed)

    segs = re.split(r'(\x00[^\x01]*\x01)', processed)
    for seg in segs:
        if seg.startswith('\x00') and seg.endswith('\x01'):
            sub_text = seg[1:-1]
            r = p.add_run(sub_text)
            r.font.subscript = True
            r.font.size = Pt(12)
        elif seg:
            add_run(p, seg, size=Pt(12))
    return p


def formula(cell, latex):
    """Numbered display equation using borderless 3-column table.

    Layout: [0.8cm spacer | 11.2cm centered OMML | 2.6cm right (N)]
    Falls back to Courier text on conversion failure.
    """
    from .utils import latex_to_mathml, mathml_to_omml

    n = _eq.next()
    try:
        mathml = latex_to_mathml(latex, 'block')
        omml = mathml_to_omml(mathml, 'block')
    except Exception:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, f'{latex}    ({n})', size=Pt(12))
        return

    table = cell.add_table(rows=1, cols=3)
    table.autofit = False
    _remove_table_borders(table)

    cell_spacer, cell_eq, cell_num = table.rows[0].cells
    cell_spacer.width = Cm(0.8)
    cell_eq.width = Cm(11.2)
    cell_num.width = Cm(2.6)

    # Inject OMML into center cell
    eq_para = cell_eq.paragraphs[0]
    eq_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_para._element.append(etree.fromstring(etree.tostring(omml)))

    # Equation number in right cell
    num_para = cell_num.paragraphs[0]
    num_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(num_para, f'({n})', size=Pt(12))

    # Vertical center alignment for all cells
    for c in table.rows[0].cells:
        tcPr = c._tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

    # Tiny gap after equation
    spacer_p = cell.add_paragraph('')
    spacer_p.paragraph_format.space_before = Pt(2)
    spacer_p.paragraph_format.space_after = Pt(4)
    spacer_p.paragraph_format.line_spacing = 0.2


def _remove_table_borders(table):
    """Remove all borders from a table — make it invisible."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{border_name}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblPr.append(borders)


def w3table(cell, headers, rows, caption=None):
    """Three-line table: top/bottom 1.5pt, header-bottom 0.75pt, no vertical lines.

    All cell text centered horizontally and vertically.
    """
    if caption:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.space_before = Pt(8)
        pf.space_after = Pt(2)
        add_run(p, caption, bold=True, size=Pt(10.5))

    table = cell.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    ))

    for ci, h in enumerate(headers):
        c2 = table.rows[0].cells[ci]
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p2.paragraph_format
        pf.line_spacing = 1.5
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        add_run(p2, h, bold=True, size=Pt(10.5))
        set_cell_border(c2, bottom={"sz": 6, "val": "single", "color": "000000"})
        tcPr = c2._tc.get_or_add_tcPr()
        va = OxmlElement('w:vAlign')
        va.set(qn('w:val'), 'center')
        tcPr.append(va)

    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            c2 = table.rows[ri + 1].cells[ci]
            p2 = c2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p2.paragraph_format
            pf.line_spacing = 1.5
            pf.space_before = Pt(1)
            pf.space_after = Pt(1)
            add_run(p2, val, size=Pt(10.5))
            tcPr = c2._tc.get_or_add_tcPr()
            va = OxmlElement('w:vAlign')
            va.set(qn('w:val'), 'center')
            tcPr.append(va)

    return table


def img(cell, path, width_inches=4.2):
    """Insert a centered image."""
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    ipath = Path(path)
    if ipath.exists():
        p.add_run().add_picture(str(ipath), width=Inches(width_inches))


def figcap(cell, text):
    """Figure caption: centered, 10.5pt."""
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(2)
    pf.space_after = Pt(6)
    add_run(p, text, size=Pt(10.5))


# ── document-level helpers ───────────────────────────────────────────────

def create_document(page_width_cm=21.0, page_height_cm=29.7,
                    top_margin_cm=2.54, bottom_margin_cm=2.54,
                    left_margin_cm=1.92, right_margin_cm=1.57,
                    font_ascii='Calibri', font_east_asia='宋体',
                    font_size_pt=12, line_spacing=1.5) -> Document:
    """Create a Document with CJK-configured Normal style."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(page_width_cm)
    section.page_height = Cm(page_height_cm)
    section.top_margin = Cm(top_margin_cm)
    section.bottom_margin = Cm(bottom_margin_cm)
    section.left_margin = Cm(left_margin_cm)
    section.right_margin = Cm(right_margin_cm)

    style = doc.styles['Normal']
    style.font.name = font_ascii
    style.font.size = Pt(font_size_pt)
    style.paragraph_format.line_spacing = line_spacing
    rPr = style.element.get_or_add_rPr()
    ex = rPr.find(qn('w:rFonts'))
    if ex is not None:
        rPr.remove(ex)
    rPr.insert(0, parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="{font_ascii}" '
        f'w:eastAsia="{font_east_asia}" w:hAnsi="{font_ascii}" w:cs="Times New Roman"/>'
    ))
    return doc


def add_content_table(doc_or_cell, rows=4, cell_width_cm=17.47):
    """Create a bordered (sz=4) N×1 table for content sections."""
    tbl = doc_or_cell.add_table(rows=rows, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_tbl_borders(tbl, sz=4)
    for row in tbl.rows:
        row.cells[0].width = Cm(cell_width_cm)
        clear_first_para(row.cells[0])
    return tbl


def reset_equation_counter():
    """Reset the global equation counter for a new document."""
    _eq.n = 0


def body_with_math(cell, text, space_before=0, size=Pt(12)):
    """Body paragraph with inline $...$ LaTeX rendered as OMML.

    Splits text on $...$ boundaries: normal text becomes runs,
    LaTeX segments become inline OMML elements inserted directly
    into the paragraph XML.
    Falls back to Courier text on conversion failure.
    """
    from .utils import latex_to_mathml, mathml_to_omml

    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(0)

    # Split on $...$ — captures content between $ signs
    segments = _re_builtin.split(r'(?<!\\)\$([^$]+)(?<!\\)\$', text)
    for i, seg in enumerate(segments):
        if seg is None or seg == '':
            continue
        if i % 2 == 1:  # LaTeX segment (odd indices in split result)
            latex = seg.strip()
            if not latex:
                continue
            try:
                mathml = latex_to_mathml(latex, 'inline')
                omml = mathml_to_omml(mathml, 'inline')
                # Insert OMML element directly into paragraph XML
                p._element.append(etree.fromstring(etree.tostring(omml)))
            except Exception:
                # Fallback: render as Courier text
                safe = _sanitize_text(latex)
                run = p.add_run(f'${safe}$')
                run.font.name = 'Courier'
                run.font.size = size
                rPr = run._element.get_or_add_rPr()
                rf = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体"/>')
                ex = rPr.find(qn('w:rFonts'))
                if ex is not None:
                    rPr.remove(ex)
                rPr.insert(0, rf)
        else:
            # Normal text segment
            if seg.strip():
                add_run(p, seg, size=size)
    return p


def _sanitize_text(text: str) -> str:
    """Remove XML-invalid control characters."""
    return _re_builtin.sub('[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
