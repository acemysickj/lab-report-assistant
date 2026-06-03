"""DOCX v2 renderer: HTML → python-docx via HTMLParser.

Adapted from html2docx (pqzx/html2docx, MIT).
Extensions: CJK, page setup, HTML5 elements, dual-track formula system.
"""
from __future__ import annotations
import io, os, re, logging
from html.parser import HTMLParser
from pathlib import Path

import docx, docx.table
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup
from lxml import etree

log = logging.getLogger(__name__)

# ── helpers ───────────────────────────────────────────────────────────

def _remove_whitespace(text: str, leading: bool = False, trailing: bool = False) -> str:
    if leading:
        text = re.sub(r'^\s*\n+\s*', '', text)
    if trailing:
        text = re.sub(r'\s*\n+\s*$', '', text)
    text = re.sub(r'\s*\n\s*', ' ', text)
    return re.sub(r'\s+', ' ', text)

def _delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def _remove_last_occurrence(ls: list, x) -> None:
    ls.pop(len(ls) - ls[::-1].index(x) - 1)

# ── style maps ────────────────────────────────────────────────────────

_FONT_STYLES = {
    'b': 'bold', 'strong': 'bold', 'th': 'bold', 'summary': 'bold',
    'em': 'italic', 'i': 'italic',
    'u': 'underline', 's': 'strike',
    'sup': 'superscript', 'sub': 'subscript',
}
_FONT_NAMES = {'code': 'Courier', 'pre': 'Courier'}
_TRANSPARENT = {'div', 'section', 'article', 'main', 'details'}

# ── LaTeX formula detection (used only on plain-text handle_data input) ──

_FORMULA_SPLIT = re.compile(
    r'(\\begin\{(?:equation|align\*?|gather\*?|cases)\}.*?'
    r'\\end\{(?:equation|align\*?|gather\*?|cases)\})'
    r'|(\$\$.+?\$\$)'
    r'|(\\\(.+?\\\))',
    re.DOTALL,
)

# ── page setup ────────────────────────────────────────────────────────

def _setup_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

def _set_cjk(run) -> None:
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def _make_border(name: str, sz_eighths: int) -> OxmlElement:
    """Create a w:top / w:bottom / w:insideH / w:insideV border element.

    sz_eighths=0 means 'none'.  1.5pt = 12, 0.75pt = 6.
    """
    el = OxmlElement(f'w:{name}')
    if sz_eighths == 0:
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
    else:
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz_eighths))
    el.set(qn('w:space'), '0')
    el.set(qn('w:color'), '000000')
    return el

# ── OMML conversion ───────────────────────────────────────────────────

def _latex_to_omml(latex: str, display: str = "inline"):
    """Convert LaTeX to OMML element via addFormula2docx pipeline.  None on failure."""
    from .utils import latex_to_mathml, mathml_to_omml
    try:
        mathml = latex_to_mathml(latex, display)
        return mathml_to_omml(mathml, display)
    except Exception as e:
        log.warning("LaTeX→OMML failed: %s  [%s]", e, latex[:60])
        return None

# ── main renderer class ───────────────────────────────────────────────

class _BaseRenderer(HTMLParser):
    """Streaming HTML → python-docx converter."""

    def __init__(self):
        super().__init__()
        self.table_row_selectors = [
            'table > tr', 'table > thead > tr',
            'table > tbody > tr', 'table > tfoot > tr',
        ]
        self._base_path = ''

    # ── setup ──────────────────────────────────────────────────────

    def _init_state(self, document=None):
        self.tags: dict[str, object] = {}
        self._list_stack: list[str] = []
        self._in_figcaption = False

        # ── clean state: ignored tags ──
        self._ignored_tag: str | None = None       # head / script / style

        # ── clean state: formula capture ──
        self._formula_capture_latex: str | None = None
        self._formula_capture_display: str = "inline"
        self._formula_capture_html: str | None = None   # fallback visible HTML
        self._formula_capture_tag: str | None = None     # 'div' or 'span'
        self._formula_capture_depth: int = 0
        # eq-num: 暂不支持 (eq-num 文本在 <span> 内部，而捕获模式下跳过了子元素)
        # 实际报告中 div.formula 的可视 HTML 本身就包含了 eq-num 的文字渲染

        self._degraded_formulas: int = 0

        # ── clean state: equation numbering ──
        self._equation_counter: int = 0
        self._equation_numbering_enabled: bool = True  # 固定显示；后续加开关

        if document is not None:
            self.doc = document
        else:
            self.doc = Document()
            _setup_page(self.doc)
        self.document = self.doc
        self.paragraph = None
        self._table_no = 0
        self._tables: list = []

    # ── helpers ────────────────────────────────────────────────────

    @staticmethod
    def _is_formula_block(tag: str, attrs: dict) -> bool:
        classes = (attrs.get('class', '') or '').split()
        return tag == 'div' and 'formula' in classes and 'data-latex' in attrs

    @staticmethod
    def _is_formula_inline(tag: str, attrs: dict) -> bool:
        classes = (attrs.get('class', '') or '').split()
        return tag == 'span' and 'math-inline' in classes and 'data-latex' in attrs

    def _is_inside_formula_capture(self) -> bool:
        return self._formula_capture_tag is not None

    # ── start-tag ──────────────────────────────────────────────────

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        current_attrs = dict(attrs)

        # ── ignored-tag guard ─────────────────────────────────
        if self._ignored_tag:
            return

        # ── formula-capture guard ──────────────────────────────
        if self._is_inside_formula_capture():
            if tag == self._formula_capture_tag:
                self._formula_capture_depth += 1
            return

        # ── detect ignored tags ────────────────────────────────
        if tag in ('head', 'script', 'style'):
            self._ignored_tag = tag
            return

        # ── detect formula blocks ──────────────────────────────
        if self._is_formula_block(tag, current_attrs):
            self._formula_capture_latex = current_attrs['data-latex']
            self._formula_capture_display = 'block'
            fml_id = current_attrs.get('data-fml-id', '')
            self._formula_capture_html = self._formula_fallbacks.get(fml_id, '')
            self._formula_capture_tag = tag
            self._formula_capture_depth = 1
            self.tags[tag] = current_attrs
            return

        if self._is_formula_inline(tag, current_attrs):
            self._formula_capture_latex = current_attrs['data-latex']
            self._formula_capture_display = 'inline'
            fml_id = current_attrs.get('data-fml-id', '')
            self._formula_capture_html = self._formula_fallbacks.get(fml_id, '')
            self._formula_capture_tag = tag
            self._formula_capture_depth = 1
            self.tags[tag] = current_attrs
            return

        # ── standard elements ──────────────────────────────────
        if tag in ('body',):
            return
        if tag in ('ol', 'ul'):
            self._list_stack.append(tag)
            return
        if tag == 'br':
            if self.paragraph is None:
                self.paragraph = self.doc.add_paragraph()
            self.paragraph.add_run().add_break()
            return

        if tag in ('p', 'pre'):
            self.paragraph = self.doc.add_paragraph()
            _set_cjk(self.paragraph.add_run(''))
            self.paragraph.clear()
        elif tag == 'li':
            self._handle_li()
        elif re.match(r'h[1-9]', tag):
            level = min(int(tag[1]), 9)
            if isinstance(self.doc, docx.document.Document):
                self.paragraph = self.doc.add_heading(level=level)
            else:
                self.paragraph = self.doc.add_paragraph()
        elif tag == 'img':
            self._handle_img(current_attrs)
            return
        elif tag == 'table':
            self._handle_table()
            return
        elif tag == 'figcaption':
            self._in_figcaption = True
            self.paragraph = self.doc.add_paragraph()
            self.paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif tag == 'hr':
            self._handle_hr()

        self.tags[tag] = current_attrs

    # ── end-tag ────────────────────────────────────────────────────

    def handle_endtag(self, tag: str) -> None:
        # ── ignored-tag guard ─────────────────────────────────
        if self._ignored_tag:
            if tag == self._ignored_tag:
                self._ignored_tag = None
            return

        # ── formula-capture guard ──────────────────────────────
        if self._is_inside_formula_capture():
            if tag == self._formula_capture_tag:
                self._formula_capture_depth -= 1
                if self._formula_capture_depth > 0:
                    return
                self._finish_formula_capture()
                self.tags.pop(tag, None)
            return

        if tag in ('ol', 'ul') and self._list_stack:
            _remove_last_occurrence(self._list_stack, tag)
            return
        if tag == 'table':
            self._table_no += 1
            self.doc = self.document
            self.paragraph = None
            return
        if tag == 'figcaption':
            self._in_figcaption = False
            self.paragraph = None
            return
        if tag in _TRANSPARENT:
            self.paragraph = None
        self.tags.pop(tag, None)

    # ── text data ──────────────────────────────────────────────────

    def handle_data(self, data: str) -> None:
        if self._ignored_tag or self._is_inside_formula_capture():
            return

        if 'pre' not in self.tags:
            data = _remove_whitespace(data, leading=True, trailing=True)
        if not data:
            return

        if self.paragraph is None:
            self.paragraph = self.doc.add_paragraph()

        if 'pre' in self.tags or 'code' in self.tags:
            self._add_text_run(data)
            return

        self._render_text_with_formulas(data)

    # ── formula capture lifecycle ──────────────────────────────────

    def _finish_formula_capture(self):
        """Convert captured data-latex to OMML.  On failure, render visible HTML."""
        latex = self._formula_capture_latex or ''
        display = self._formula_capture_display
        omml = _latex_to_omml(latex, display)

        if omml is not None:
            self._emit_omml_formula(omml, display)
        else:
            self._degraded_formulas += 1
            # Fallback: render the visible HTML from inside the formula element.
            # The original HTML (sub/sup/entities) is preserved because the
            # captured tag's inner content was never fed to the parser.
            # We re-parse the visible HTML as a fresh renderer so sub/sup are kept.
            if self._formula_capture_html:
                child = _BaseRenderer()
                child._base_path = self._base_path
                child._init_state(document=self.doc)
                child._run(self._formula_capture_html)
                # child may have created paragraphs — close current paragraph context
                self.paragraph = None
            else:
                # No HTML fallback available — show raw LaTeX in monospace
                if display == 'block':
                    self.paragraph = self.doc.add_paragraph()
                    self.paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = (self.paragraph or self.doc.add_paragraph()).add_run(f'[{latex}]')
                run.font.name = 'Courier'
                self.paragraph = None if display == 'block' else self.paragraph

        self._formula_capture_latex = None
        self._formula_capture_html = None
        self._formula_capture_tag = None
        self._formula_capture_depth = 0

    def _emit_omml_formula(self, omml, display: str):
        """Insert an already-converted OMML element into the document."""
        if display == 'block':
            if self._equation_numbering_enabled:
                self._emit_numbered_equation(omml)
            else:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p._element.append(etree.fromstring(etree.tostring(omml)))
                self.paragraph = None
        else:
            if self.paragraph is None:
                self.paragraph = self.doc.add_paragraph()
            self.paragraph._element.append(etree.fromstring(etree.tostring(omml)))

    # ── equation numbering helpers ───────────────────────────────────

    def _next_equation_number(self) -> int:
        """Return the next sequential equation number."""
        self._equation_counter += 1
        return self._equation_counter

    @staticmethod
    def _remove_table_borders(table):
        """Remove all borders from a table (make it invisible)."""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        for old in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(old)

        borders = OxmlElement('w:tblBorders')
        for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            borders.append(border)
        tblPr.append(borders)

    def _emit_numbered_equation(self, omml):
        """Create a borderless 3-column table: spacer | centered-equation | (number)."""
        table = self.doc.add_table(rows=1, cols=3)
        table.autofit = False

        cell_spacer, cell_eq, cell_num = table.rows[0].cells
        cell_spacer.width = Cm(0.8)
        cell_eq.width = Cm(11.2)
        cell_num.width = Cm(2.6)

        # Remove all borders
        self._remove_table_borders(table)

        # Inject OMML into center cell's first paragraph
        eq_para = cell_eq.paragraphs[0]
        eq_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq_para._element.append(etree.fromstring(etree.tostring(omml)))

        # Add equation number to right cell
        num_para = cell_num.paragraphs[0]
        num_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        eq_num = self._next_equation_number()
        run = num_para.add_run(f"({eq_num})")
        run.font.size = Pt(12)
        _set_cjk(run)

        # Set cell vertical alignment to center
        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

        # Tiny gap after equation table
        spacer = self.doc.add_paragraph('')
        spacer.paragraph_format.space_before = Pt(2)
        spacer.paragraph_format.space_after = Pt(4)
        spacer.paragraph_format.line_spacing = 0.2

        self.paragraph = None

    # ── bare-LaTeX text splitting ──────────────────────────────────

    def _render_text_with_formulas(self, text: str):
        """Split text on LaTeX delimiters, emitting runs and OMML in order."""
        parts = _FORMULA_SPLIT.split(text)
        for part in (p for p in parts if p):
            if part.startswith('\\begin{') or part.startswith('$$'):
                self._emit_block_latex(part.strip())
            elif part.startswith('\\('):
                self._emit_inline_latex(part[2:-2].strip())
            else:
                if part.strip():
                    self._add_text_run(part)

    def _emit_block_latex(self, raw: str):
        """Emit a block formula from bare LaTeX ($$...$$ or \\begin{...}\\end{...})."""
        latex = raw
        if raw.startswith('$$') and raw.endswith('$$'):
            latex = raw[2:-2].strip()
        omml = _latex_to_omml(latex, 'block')
        if omml is not None:
            self._emit_omml_formula(omml, 'block')
        else:
            self._degraded_formulas += 1
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'[{latex}]')
            run.font.name = 'Courier'
        self.paragraph = None  # don't reuse — let next block element create its own

    def _emit_inline_latex(self, latex: str):
        """Emit an inline OMML formula from bare \\(...\\)."""
        omml = _latex_to_omml(latex, 'inline')
        if omml is not None:
            if self.paragraph is None:
                self.paragraph = self.doc.add_paragraph()
            self.paragraph._element.append(etree.fromstring(etree.tostring(omml)))
        else:
            self._degraded_formulas += 1
            run = (self.paragraph or self.doc.add_paragraph()).add_run(f'[{latex}]')
            run.font.name = 'Courier'

    # ── run helpers ────────────────────────────────────────────────

    def _add_text_run(self, text: str):
        run = self.paragraph.add_run(text)
        _set_cjk(run)
        if self._in_figcaption:
            run.font.size = Pt(10)
        self._apply_inline_styles(run)

    def _apply_inline_styles(self, run) -> None:
        for tag_name in self.tags:
            if tag_name in _FONT_STYLES:
                setattr(run.font, _FONT_STYLES[tag_name], True)
            if tag_name in _FONT_NAMES:
                run.font.name = _FONT_NAMES[tag_name]

    # ── lists ──────────────────────────────────────────────────────

    def _handle_li(self):
        list_depth = len(self._list_stack)
        list_type = self._list_stack[-1] if list_depth else 'ul'
        style = 'List Number' if list_type == 'ol' else 'List Bullet'
        self.paragraph = self.doc.add_paragraph(style=style)
        self.paragraph.paragraph_format.left_indent = Inches(min(list_depth * 0.5, 5.5))
        self.paragraph.paragraph_format.line_spacing = 1

    # ── images ─────────────────────────────────────────────────────

    def _handle_img(self, attrs: dict):
        src = attrs.get('src', '')
        if not src:
            self.paragraph = self.doc.add_paragraph('[图片: 无路径]')
            return
        if self._base_path and not os.path.isabs(src):
            src = os.path.join(self._base_path, src)
        src = os.path.normpath(src)
        try:
            if isinstance(self.doc, docx.document.Document):
                self.doc.add_picture(src, width=Inches(5.5))
                self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(src, width=Inches(5.5))
        except Exception as e:
            log.warning('Image insert failed: %s [%s]', e, Path(src).name)
            fallback_label = attrs.get('alt', '') or Path(src).name
            self.doc.add_paragraph(f'[图片: {fallback_label}]')

    # ── hr ─────────────────────────────────────────────────────────

    def _handle_hr(self):
        self.paragraph = self.doc.add_paragraph()
        pPr = self.paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── tables ─────────────────────────────────────────────────────

    def _handle_table(self):
        if self._table_no >= len(self._tables):
            return
        table_soup = self._tables[self._table_no]
        rows, cols = self._get_table_dimensions(table_soup)
        table = self.doc.add_table(rows=rows, cols=cols)

        for i, tr in enumerate(self._get_table_rows(table_soup)):
            cells = tr.find_all(['th', 'td'], recursive=False)
            for j, cell in enumerate(cells):
                if j >= cols:
                    continue
                self._render_cell(cell, table.cell(i, j))

        self._apply_three_line_table(table)
        self._center_cell_paragraphs(table)
        # Tiny gap after table
        spacer = self.doc.add_paragraph('')
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(4)
        spacer.paragraph_format.line_spacing = 0.2

    # ── three-line table ───────────────────────────────────────────

    @staticmethod
    def _apply_three_line_table(table):
        """Apply strict three-line table style via OXML.

        Top 1.5pt, header-bottom 0.75pt, bottom 1.5pt.
        No vertical lines, no internal horizontal lines.
        Cells centered with tight padding.
        """
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Remove any existing borders
        for old in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(old)

        borders = OxmlElement('w:tblBorders')
        borders.append(_make_border('top', 12))       # 1.5 pt
        borders.append(_make_border('bottom', 12))    # 1.5 pt
        borders.append(_make_border('left', 0))       # none
        borders.append(_make_border('right', 0))      # none
        borders.append(_make_border('insideH', 0))    # none
        borders.append(_make_border('insideV', 0))    # none
        tblPr.append(borders)

        rows = tbl.findall(qn('w:tr'))
        for i, tr in enumerate(rows):
            for tc in tr.findall(qn('w:tc')):
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)

                # Center cell text vertically
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)

                # Header row: bottom border 0.75pt
                if i == 0:
                    for old_b in tcPr.findall(qn('w:tcBorders')):
                        tcPr.remove(old_b)
                    tcBorders = OxmlElement('w:tcBorders')
                    tcBorders.append(_make_border('bottom', 6))  # 0.75 pt
                    tcPr.append(tcBorders)

    @staticmethod
    def _center_cell_paragraphs(table):
        """Center-align all paragraphs in all cells."""
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _render_cell(self, soup_cell, docx_cell):
        if docx_cell.paragraphs and docx_cell.paragraphs[0].text == '':
            _delete_paragraph(docx_cell.paragraphs[0])
        child = _BaseRenderer()
        child._base_path = self._base_path
        child._init_state(document=docx_cell)
        html = str(soup_cell)
        if soup_cell.name == 'th':
            html = f'<b>{html}</b>'
        child._run(html)
        if not child.doc.paragraphs:
            child.doc.add_paragraph('')
        # Tighten cell paragraph spacing
        for para in docx_cell.paragraphs:
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)

    @staticmethod
    def _get_table_dimensions(table_soup):
        selectors = ['table > tr', 'table > thead > tr',
                     'table > tbody > tr', 'table > tfoot > tr']
        rows = table_soup.select(', '.join(selectors), recursive=False)
        if not rows:
            return 0, 0
        cols = rows[0].find_all(['th', 'td'], recursive=False)
        return len(rows), len(cols)

    def _get_table_rows(self, table_soup):
        return table_soup.select(', '.join(self.table_row_selectors), recursive=False)

    def _extract_tables(self, soup):
        all_tables = soup.find_all('table')
        filtered = []
        skip = 0
        for t in all_tables:
            if skip:
                skip -= 1
                continue
            filtered.append(t)
            skip = len(t.find_all('table'))
        self._tables = filtered
        self._table_no = 0

    # ── entry points ───────────────────────────────────────────────

    def _run(self, html: str):
        soup = BeautifulSoup(html, 'html.parser')
        self._extract_tables(soup)
        # Pre-process formula blocks: inject data-fml-id for fallback HTML lookup
        self._formula_fallbacks: dict[str, str] = {}
        _fml_counter = 0
        for tag_name, cls_attr in [('div', 'formula'), ('span', 'math-inline')]:
            for el in soup.find_all(tag_name, class_=cls_attr):
                if el.get('data-latex'):
                    fml_id = f'fml_{_fml_counter}'
                    _fml_counter += 1
                    el['data-fml-id'] = fml_id
                    # Store inner HTML (children rendered as string) for fallback
                    inner = ''.join(str(c) for c in el.contents)
                    self._formula_fallbacks[fml_id] = inner
        self.feed(str(soup))

    def add_html_to_document(self, html: str, document):
        self._init_state(document)
        self._run(html)

    def add_html_to_cell(self, html: str, cell):
        self._init_state(cell)
        self._run(html)


# ── public API ────────────────────────────────────────────────────────

def render_html_to_docx(html: str, base_path: str = '') -> bytes:
    doc = Document()
    _setup_page(doc)
    renderer = _BaseRenderer()
    renderer._base_path = base_path
    renderer.add_html_to_document(html, doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
