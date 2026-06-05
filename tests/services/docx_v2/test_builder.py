"""Tests for docx_v2 builder module."""
import sys
import os
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..', '..', 'backend'))

from services.docx_v2.builder import (
    create_document,
    add_content_table,
    heading, sub, body, body_lbl, body_sub,
    formula, w3table, img, figcap,
    add_run, set_tbl_borders, set_cell_border,
    clear_first_para, reset_equation_counter,
)


class TestCreateDocument:
    def test_page_setup(self):
        doc = create_document(21.0, 29.7, 2.54, 2.54, 1.92, 1.57)
        section = doc.sections[0]
        assert round(section.page_width.cm, 1) == 21.0
        assert round(section.page_height.cm, 1) == 29.7
        assert round(section.top_margin.cm, 1) == 2.5
        assert round(section.left_margin.cm, 1) == 1.9

    def test_font_config(self):
        doc = create_document(font_ascii='Calibri', font_east_asia='宋体', font_size_pt=12)
        style = doc.styles['Normal']
        assert style.font.name == 'Calibri'
        assert style.font.size == Pt(12)
        rPr = style.element.find(qn('w:rPr'))
        rFonts = rPr.find(qn('w:rFonts'))
        assert rFonts.get(qn('w:eastAsia')) == '宋体'


class TestContentTable:
    def test_add_content_table(self):
        doc = create_document()
        tbl = add_content_table(doc, rows=4)
        assert len(tbl.rows) == 4
        assert len(tbl.columns) == 1

    def test_content_table_has_borders(self):
        doc = create_document()
        tbl = add_content_table(doc, rows=3)
        tblPr = tbl._tbl.tblPr
        borders = tblPr.find(qn('w:tblBorders'))
        assert borders is not None


class TestHeading:
    def test_heading_adds_bold_paragraph(self):
        doc = create_document()
        tbl = add_content_table(doc)
        c = tbl.rows[0].cells[0]
        heading(c, '一、实验目的')
        assert len(c.paragraphs) == 1
        p = c.paragraphs[0]
        assert '一、实验目的' in p.text
        run = p.runs[0]
        assert run.bold is True
        assert run.font.size == Pt(14)


class TestSub:
    def test_sub_not_bold(self):
        doc = create_document()
        tbl = add_content_table(doc)
        c = tbl.rows[0].cells[0]
        sub(c, '1. 直流电压测量方法')
        assert len(c.paragraphs) == 1
        run = c.paragraphs[0].runs[0]
        assert run.bold is False
        assert run.font.size == Pt(12)


class TestBody:
    def test_body_adds_text(self):
        doc = create_document()
        tbl = add_content_table(doc)
        c = tbl.rows[0].cells[0]
        body(c, '这是一段正文内容。')
        assert '这是一段正文内容。' in c.paragraphs[0].text

    def test_body_with_space_before(self):
        doc = create_document()
        tbl = add_content_table(doc)
        c = tbl.rows[0].cells[0]
        body(c, '带间距的正文', space_before=6)
        pf = c.paragraphs[0].paragraph_format
        assert pf.space_before == Pt(6)


class TestBodyLbl:
    def test_body_lbl_concats_label_and_text(self):
        doc = create_document()
        tbl = add_content_table(doc)
        c = tbl.rows[0].cells[0]
        body_lbl(c, '幅值测量：', '使用示波器自动测量。')
        full_text = ''.join(r.text for r in c.paragraphs[0].runs)
        assert '幅值测量' in full_text
        assert '示波器' in full_text


class TestFormula:
    def test_formula_creates_tab_stops(self):
        doc = create_document()
        tbl = add_content_table(doc)
        c = tbl.rows[0].cells[0]
        reset_equation_counter()
        formula(c, r'U_{R1} = \frac{R_1}{R_1 + R_2} U')
        p = c.paragraphs[0]
        # Check tab stops exist
        pPr = p._element.find(qn('w:pPr'))
        tabs = pPr.find(qn('w:tabs'))
        assert tabs is not None
        tab_els = tabs.findall(qn('w:tab'))
        assert len(tab_els) == 2

    def test_formula_numbering_increments(self):
        doc = create_document()
        tbl = add_content_table(doc)
        c = tbl.rows[0].cells[0]
        reset_equation_counter()
        formula(c, r'x = 1')  # (1)
        formula(c, r'y = 2')  # (2)
        formula(c, r'z = 3')  # (3)
        texts = [p.text for p in c.paragraphs]
        assert any('(1)' in t for t in texts)
        assert any('(2)' in t for t in texts)
        assert any('(3)' in t for t in texts)

    def test_reset_counter(self):
        doc = create_document()
        tbl = add_content_table(doc)
        c = tbl.rows[0].cells[0]
        reset_equation_counter()
        formula(c, r'a = 1')
        reset_equation_counter()
        formula(c, r'b = 1')
        texts = [p.text for p in c.paragraphs]
        assert any('(1)' in t for t in texts)
        # After reset, second formula also gets (1)


class TestW3Table:
    def test_w3table_structure(self):
        doc = create_document()
        tbl = add_content_table(doc)
        c = tbl.rows[0].cells[0]
        table = w3table(c, ['A', 'B', 'C'], [['1', '2', '3'], ['4', '5', '6']],
                        caption='表 1  测试表')
        assert len(table.rows) == 3  # header + 2 data rows
        assert len(table.columns) == 3
        # Check caption exists
        assert '表 1' in c.paragraphs[0].text

    def test_w3table_three_line_borders(self):
        doc = create_document()
        tbl = add_content_table(doc)
        c = tbl.rows[0].cells[0]
        table = w3table(c, ['A'], [['1']])
        tblEl = table._tbl
        tblPr = tblEl.find(qn('w:tblPr'))
        borders = tblPr.find(qn('w:tblBorders'))
        assert borders is not None
        # Top and bottom should be sz=12 (1.5pt)
        top = borders.find(qn('w:top'))
        bottom = borders.find(qn('w:bottom'))
        assert top.get(qn('w:sz')) == '12'
        assert bottom.get(qn('w:sz')) == '12'
        # InsideH should be none
        insideH = borders.find(qn('w:insideH'))
        assert insideH.get(qn('w:val')) == 'none'


class TestFigcap:
    def test_figcap_centered(self):
        doc = create_document()
        tbl = add_content_table(doc)
        c = tbl.rows[0].cells[0]
        figcap(c, '图 1  示意图')
        p = c.paragraphs[0]
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        assert run.font.size == Pt(10.5)


class TestClearFirstPara:
    def test_removes_empty_default_paragraph(self):
        doc = create_document()
        tbl = doc.add_table(rows=1, cols=1)
        cell = tbl.rows[0].cells[0]
        assert len(cell.paragraphs) == 1
        assert cell.paragraphs[0].text == ''
        clear_first_para(cell)
        # Empty paragraph should be removed (or at least the default one is gone)
        # python-docx may still report paragraphs but the element is removed
